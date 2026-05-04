from langchain_core.documents import Document
from pypdf import PdfReader
import docx
import os

def read_txt(path):
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def read_pdf(path):
    reader = PdfReader(path)
    text = ""

    for page in reader.pages:
        text += page.extract_text() or ""

    return text


def read_docx(path):
    doc = docx.Document(path)
    return "\n".join([para.text for para in doc.paragraphs])


def process_uploaded_file(file_path):
    ext = os.path.splitext(file_path)[-1].lower()

    if ext == ".txt":
        content = read_txt(file_path)

    elif ext == ".pdf":
        content = read_pdf(file_path)

    elif ext == ".docx":
        content = read_docx(file_path)

    else:
        raise ValueError("Unsupported file type")

    return [
        Document(
            page_content=content,
            metadata={"source": file_path, "type": ext}
        )
    ]